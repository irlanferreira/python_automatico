from docx import Document
from datetime import datetime

documento = Document('relatorio_inicial.docx')

documento.paragraphs[0].text = "Relatório Financeiro - Revisado"
documento.paragraphs[1].add_run(" Os valores estão sujeitos a conferência.")

tabela = documento.tables[0]
tabela.cell(2, 2).text = "13000"

linha = tabela.add_row().cells
linha[0].text = 'Abril'
linha[1].text = '13000'
linha[2].text = '9000'

agora = datetime.now()
documento.add_paragraph(f"Relatório revisado em Python em {agora.strftime("%d/%m/%Y")}")

documento.save("relatorio_editado.docx")