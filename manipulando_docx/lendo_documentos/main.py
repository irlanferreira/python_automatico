from docx import Document

documento = Document('exemplo.docx')
# for p in documento.paragraphs:
#     print(p.text)

# for tabela in documento.tables:
#     for linha in tabela.rows:
#         dados = [celula.text for celula in linha.cells]
#         print(dados)

# titulo = documento.paragraphs[0]
# titulo.text = "Título do Documento - Editado"

tabela = documento.tables[0]
tabela.cell(1, 2).text = "3"

documento.add_paragraph("Este e um novo paragrafo!")

documento.save('exemplo.docx')