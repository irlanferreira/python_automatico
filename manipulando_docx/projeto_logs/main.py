import re
from datetime import datetime
from docx import Document

def extrair_dados(log):
    logs = []
    expressao = r"(\d{4}-\d{2}-\d{2})\s*(\d{2}:\d{2}:\d{2})\s*(\w+)\s*(.+)"
    resultados = re.findall(expressao, log)
    if resultados:
        for resultado in resultados:
            logs.append(
                {
                    'data':resultado[0],
                    'hora':resultado[1],
                    'tipo':resultado[2],
                    'mensagem':resultado[3]
                }
            )
        return logs
    else:
        print("Não encontrei nada!")
        return False

def criar_relatorio():
    with open('logs.txt', encoding='utf-8') as log:
        log_completo = log.read()
    logs = extrair_dados(log_completo)
    
    documento = Document()
    documento.add_heading('Relatório de Análise de Logs')

    logs_tipos_qtd = {}

    for log in logs:
        if log['tipo'] in logs_tipos_qtd:
            logs_tipos_qtd[log['tipo']] += 1
        else:
            logs_tipos_qtd[log['tipo']] = 1
    
    documento.add_paragraph(f"Total de ocorrências para cada tipo:")
    for tipo, qtd in logs_tipos_qtd.items():
        documento.add_paragraph(f"{tipo}: {qtd}")
    
    documento.add_paragraph("Erros ocorridos:")
    for log in logs:
        if log['tipo'] == "ERROR":
            documento.add_paragraph(f"{log['data']} {log['hora']}: {log['mensagem']}")

    documento.add_page_break()

    registros_por_dia = {}
    for log in logs:
        if log['data'] not in registros_por_dia:
            registros_por_dia[log['data']] = {}
        if log['tipo'] not in registros_por_dia[log['data']]:
            registros_por_dia[log['data']][log['tipo']] = 0
        registros_por_dia[log['data']][log['tipo']] += 1
    
    tabela = documento.add_table(0, 5, 'Light List Accent 1')
    linha = tabela.add_row().cells
    linha[0].text = 'Data'
    linha[1].text = 'INFO'
    linha[2].text = 'ERROR'
    linha[3].text = 'WARNING'
    linha[4].text = 'DEBUG'

    for data in sorted(registros_por_dia):
        linha = tabela.add_row().cells
        linha[0].text = data
        linha[1].text = str(registros_por_dia[data].get('INFO', 0))
        linha[2].text = str(registros_por_dia[data].get('ERROR', 0))
        linha[3].text = str(registros_por_dia[data].get('WARNING', 0))
        linha[4].text = str(registros_por_dia[data].get('DEBUG', 0))

    documento.save("relatorio_teste.docx")


criar_relatorio()