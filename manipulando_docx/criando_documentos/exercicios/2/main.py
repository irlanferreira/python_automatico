from docx import Document

documento = Document()

titulo = documento.add_heading("Lista de Funcionários")

tabela = documento.add_table(0, 3, "Light List Accent 1")

cabecalho_tabela = tabela.add_row().cells
cabecalho_tabela[0].text = "Nome"
cabecalho_tabela[1].text = "Cargo"
cabecalho_tabela[2].text = "Salário"

funcionarios = [
    ('Ana', 'Programadora', "5600"),
    ('Pedro', 'Marketero', "3500"),
    ("João", 'Porteiro', "2000")
]

for funcionario in funcionarios:
    linha = tabela.add_row().cells
    linha[0].text = funcionario[0]
    linha[1].text = funcionario[1]
    linha[2].text = funcionario[2]

documento.save('funcionarios.docx')