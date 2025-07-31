from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

documento = Document()
cabecalho = documento.add_heading("Contrato de Prestação de Serviço", level=1)
cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER

documento.add_paragraph("Este contrato tem como objeto a prestação de serviços de desenvolvimento de software.")

documento.add_paragraph("Fornecer as informações necessárias", "List Bullet")
documento.add_paragraph("Realizar os pagamentos no prazo", "List Bullet")
documento.add_paragraph("Disponibilizar os recursos necessários", "List Bullet")