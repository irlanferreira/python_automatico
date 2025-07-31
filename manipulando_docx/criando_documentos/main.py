from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

documento = Document()

cabecalho = documento.add_heading("", level=1)
cabecalho.add_run("Relatório Automático").font.size = Pt(16)
cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER

documento.add_paragraph("").add_run("Este texto foi gerado com Python e python-docx.").font.bold = True

documento.add_paragraph("Primeiro Item", style="List Bullet")
documento.add_paragraph("Segundo Item", style="List Bullet")
documento.add_paragraph("Terceiro Item", style="List Bullet")

documento.add_paragraph("Primeiro Item", style="ListNumber")
documento.add_paragraph("Segundo Item", style="ListNumber")
documento.add_paragraph("Terceiro Item", style="ListNumber")

# documento.add_page_break()
# documento.add_paragraph('Teste')

tabela = documento.add_table(0, 3)
tabela.style = "Light List Accent 1"

cabecalho_tabela = tabela.add_row().cells
cabecalho_tabela[0].text = "Produto"
cabecalho_tabela[1].text = 'Preço'
cabecalho_tabela[2].text = 'Quantidade'

linha = tabela.add_row().cells
linha[0].text = 'Mouse'
linha[1].text = "14"
linha[2].text = '4'

documento.add_picture("gato.png", Inches(2))
documento.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


documento.save('relatorio.docx')